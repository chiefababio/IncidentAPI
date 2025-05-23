from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from typing import List
from fastapi.responses import FileResponse
import os
from openai import OpenAI
from docx import Document
from fpdf import FPDF
from dotenv import load_dotenv

load_dotenv()

app = FastAPI()
client = OpenAI()

class Incident(BaseModel):
    title: str
    description: str
    location: str
    date: str

incidents = []
id_counter = 1

@app.post("/api/incidents")
async def create_incident(incident: Incident):
    global id_counter
    incident_data = incident.dict()
    incident_data['id'] = id_counter
    id_counter += 1

    # Use ChatGPT for root cause and corrective action
    prompt = f"Analyze this safety incident:\nTitle: {incident.title}\nDescription: {incident.description}\nLocation: {incident.location}\nDate: {incident.date}\nProvide root causes, underlying causes, and suggest corrective actions."
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=500
    )
    analysis = response.choices[0].message.content

    incident_data['analysis'] = analysis
    incidents.append(incident_data)
    return {"incident": incident_data}

@app.get("/api/incidents/{incident_id}/report/{report_type}")
async def generate_report(incident_id: int, report_type: str):
    incident = next((i for i in incidents if i["id"] == incident_id), None)
    if not incident:
        raise HTTPException(status_code=404, detail="Incident not found")

    if report_type == 'pdf':
        filename = f"incident_{incident_id}.pdf"
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt=f"Incident Report: {incident['title']}", ln=True, align='C')
        pdf.multi_cell(0, 10, txt=f"Description: {incident['description']}")
        pdf.multi_cell(0, 10, txt=f"Location: {incident['location']}")
        pdf.multi_cell(0, 10, txt=f"Date: {incident['date']}")
        pdf.multi_cell(0, 10, txt=f"Analysis: {incident['analysis']}")
        pdf.output(filename)
        return FileResponse(filename, media_type='application/pdf', filename=filename)

    elif report_type == 'docx':
        filename = f"incident_{incident_id}.docx"
        doc = Document()
        doc.add_heading(f"Incident Report: {incident['title']}", 0)
        doc.add_paragraph(f"Description: {incident['description']}")
        doc.add_paragraph(f"Location: {incident['location']}")
        doc.add_paragraph(f"Date: {incident['date']}")
        doc.add_paragraph("Analysis:")
        doc.add_paragraph(incident['analysis'])
        doc.save(filename)
        return FileResponse(filename, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=filename)
    else:
        raise HTTPException(status_code=400, detail="Invalid report type")
